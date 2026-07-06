"""
Supporting Information — LabBook Module (Streamlit)
=====================================================
Modular app to compile Supporting Information for organic synthesis compounds.
- Save/Load session via JSON (allows partial completion across sessions)
- Raw file parsers: FTIR (.txt), TGA (NETZSCH .txt), DRX (Rigaku .txt)
- Export to Word (.docx, ACS-style)
- One entry per compound

Deploy: streamlit run app.py
Requirements: streamlit, pandas, numpy, plotly, scipy, python-docx
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from scipy.signal import find_peaks
from io import BytesIO
import json
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────────────────────────────────────
# Page config
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="Supporting Information — LabBook", layout="wide")

# ─────────────────────────────────────────────────────────────────────────────
# Session state initialization
# All persistent data lives in st.session_state.data (a single dict → JSON-friendly)
# ─────────────────────────────────────────────────────────────────────────────
NMR_TECHNIQUES = ["¹H NMR", "¹³C NMR", "DEPT-135", "DEPT-90",
                  "COSY", "HMBC", "HSQC", "NOESY", "ROESY"]
MULTIPLICITIES = ["s", "d", "t", "q", "dd", "ddd", "dt", "td", "m", "br s", "br d"]
C_TYPES = ["C", "CH", "CH₂", "CH₃"]


def default_data():
    """Return the empty data structure for a new compound entry."""
    return {
        # 1. Identification (would be pulled from Reaction Setup in full LabBook)
        "id": {
            "compound_name": "",
            "compound_code": "",
            "reaction_ref": "",
            "analyst": "",
            "date": "",
            "smiles": "",
        },
        # 2. NMR — per-technique dict; each holds metadata + table rows
        "nmr_checked": {t: False for t in NMR_TECHNIQUES},
        "nmr": {t: {"solvent": "", "freq": "", "temp": "", "rows": [], "notes": "", "pdf": ""}
                for t in NMR_TECHNIQUES},
        # 3. IR
        "ir": {"technique": "ATR", "bands": "", "notes": "", "pdf": "",
               "detected_bands": []},
        # 4. GC
        "gc": {"column": "", "temp": "", "carrier": "", "rt": "", "notes": "",
               "pdf": "", "peaks": []},
        # 5. GC-MS
        "gcms": {"column": "", "temp": "", "rt": "", "ions": "", "notes": "", "pdf": ""},
        # 6. HPLC
        "hplc": {"column": "", "mobile": "", "uv": "", "rt": "", "ee": "",
                 "notes": "", "pdf": "", "peaks": []},
        # 7. HRMS
        "hrms": {"technique": "ESI+", "formula": "", "calc": "", "found": "",
                 "notes": "", "pdf": ""},
        # 8. Physical properties
        "mp": {"value": "", "lit": ""},
        "optrot": {"alpha": "", "conc": "", "solvent": "", "temp": "25"},
        "ri": {"value": "", "temp": "20"},
        # 9. Elemental analysis
        "ea": {"rows": [{"elem": "C", "calc": "", "found": ""},
                        {"elem": "H", "calc": "", "found": ""},
                        {"elem": "N", "calc": "", "found": ""}],
               "notes": ""},
        # 10. Thermal / XRD
        "dsc": {"onset": "", "peak": "", "enthalpy": "", "notes": "", "pdf": ""},
        "tga": {"onset": "", "notes": "", "pdf": "", "detected_onset": ""},
        "xrd": {"notes": "", "pdf": "", "wavelength": "1.5406", "detected_peaks": []},
        # 11. Additional
        "extra": {"notes": "", "pdf": ""},
    }


if "data" not in st.session_state:
    st.session_state.data = default_data()

D = st.session_state.data  # shorthand


# ─────────────────────────────────────────────────────────────────────────────
# Raw-file parsers
# ─────────────────────────────────────────────────────────────────────────────
def parse_two_column_txt(content: str):
    """Generic parser for 2-column numeric .txt (used for FTIR).
    Skips any non-numeric header lines; auto-detects separator."""
    x, y = [], []
    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue
        # Split on tab, comma, semicolon or whitespace
        parts = re.split(r"[\t,;]+|\s+", line)
        if len(parts) < 2:
            continue
        try:
            xv = float(parts[0].replace(",", "."))
            yv = float(parts[1].replace(",", "."))
            x.append(xv)
            y.append(yv)
        except ValueError:
            # header line — skip
            continue
    return np.array(x), np.array(y)


def parse_tga_netzsch(content: str):
    """Parse NETZSCH STA .txt: '#' metadata lines, then '##Temp...' header,
    then semicolon-separated data. Returns temp (°C), mass (%), dtg (%/min)."""
    temp, mass, dtg = [], [], []
    data_started = False
    for line in content.splitlines():
        s = line.strip()
        if s.startswith("##"):        # column header line
            data_started = True
            continue
        if s.startswith("#") or not s:  # metadata / blank
            continue
        if not data_started:
            # Some exports skip the ## line — fall back to numeric detection
            if ";" not in s:
                continue
        parts = s.split(";")
        if len(parts) < 3:
            continue
        try:
            temp.append(float(parts[0]))
            mass.append(float(parts[2]))          # col 3 = Mass/%
            dtg.append(float(parts[3]) if len(parts) > 3 else np.nan)
        except (ValueError, IndexError):
            continue
    return np.array(temp), np.array(mass), np.array(dtg)


def parse_drx_rigaku(content: str):
    """Parse Rigaku Ultima .txt: metadata header block, then tab-separated
    '2theta<TAB>intensity'. Returns 2theta, intensity arrays."""
    two_theta, intensity = [], []
    for line in content.splitlines():
        s = line.strip()
        if not s:
            continue
        parts = re.split(r"[\t\s]+", s)
        if len(parts) < 2:
            continue
        try:
            tt = float(parts[0])
            inten = float(parts[1])
            # Guard against metadata like "40kV/20mA" being misread
            two_theta.append(tt)
            intensity.append(inten)
        except ValueError:
            continue
    return np.array(two_theta), np.array(intensity)


def detect_ir_bands(x, y, n=12):
    """Detect main IR absorption bands. Assumes transmittance/absorbance;
    finds the most prominent minima (transmittance) OR maxima (absorbance).
    Heuristic: if mean(y) > midpoint → transmittance (dips are bands)."""
    if len(x) == 0:
        return []
    # Detect orientation: transmittance data dips downward at bands
    is_transmittance = np.mean(y) > (np.max(y) + np.min(y)) / 2
    signal = -y if is_transmittance else y
    # Prominence scaled to signal range
    prom = (np.max(signal) - np.min(signal)) * 0.05
    peaks, props = find_peaks(signal, prominence=prom, distance=5)
    if len(peaks) == 0:
        return []
    # Sort by prominence, keep top n
    order = np.argsort(props["prominences"])[::-1][:n]
    band_x = sorted([round(float(x[peaks[i]])) for i in order], reverse=True)
    return band_x


def detect_tga_onset(temp, mass, dtg):
    """Estimate degradation onset as temperature at maximum mass-loss rate
    (most negative DTG). Returns rounded °C."""
    if len(temp) == 0:
        return ""
    if dtg is not None and len(dtg) == len(temp) and not np.all(np.isnan(dtg)):
        idx = int(np.nanargmin(dtg))  # most negative DTG = fastest loss
    else:
        deriv = np.gradient(mass, temp)
        idx = int(np.argmin(deriv))
    return str(round(float(temp[idx]), 1))


def detect_drx_peaks(two_theta, intensity, wavelength=1.5406, n=10):
    """Detect main XRD peaks; compute d-spacing via Bragg's law.
    d = λ / (2 sin θ), where θ = 2θ/2."""
    if len(two_theta) == 0:
        return []
    prom = (np.max(intensity) - np.min(intensity)) * 0.05
    peaks, props = find_peaks(intensity, prominence=prom, distance=10)
    if len(peaks) == 0:
        return []
    order = np.argsort(props["prominences"])[::-1][:n]
    result = []
    for i in sorted(order, key=lambda k: two_theta[peaks[k]]):
        tt = float(two_theta[peaks[i]])
        theta_rad = np.radians(tt / 2)
        d = wavelength / (2 * np.sin(theta_rad)) if np.sin(theta_rad) > 0 else 0
        result.append({"two_theta": round(tt, 2),
                       "intensity": round(float(intensity[peaks[i]]), 0),
                       "d_spacing": round(d, 4)})
    return result


# ─────────────────────────────────────────────────────────────────────────────
# ACS string builder for ¹H NMR
# ─────────────────────────────────────────────────────────────────────────────
def build_acs_1h(rows):
    """Build ACS-formatted ¹H NMR string from table rows.
    e.g. δ 7.25 (d, J = 8.2 Hz, 1H, H-3), ..."""
    parts = []
    for r in rows:
        if not r.get("delta"):
            continue
        seg = [r.get("mult", "")]
        if r.get("J"):
            seg.append(f"J = {r['J']} Hz")
        if r.get("integ"):
            seg.append(f"{r['integ']}H")
        if r.get("assign"):
            seg.append(r["assign"])
        parts.append(f"{r['delta']} ({', '.join([s for s in seg if s])})")
    return "δ " + ", ".join(parts) if parts else ""


def compute_peak_percentages(peaks):
    """Add area% to each peak row = area / total area × 100."""
    total = sum(float(p["area"]) for p in peaks if p.get("area"))
    for p in peaks:
        if p.get("area") and total > 0:
            p["pct"] = round(float(p["area"]) / total * 100, 2)
        else:
            p["pct"] = None
    return peaks, total


# ─────────────────────────────────────────────────────────────────────────────
# Word export
# ─────────────────────────────────────────────────────────────────────────────
def build_docx():
    """Generate an ACS-style Supporting Information .docx from current data."""
    doc = Document()
    # Base font: Times New Roman (ACS style)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    def h(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        return p

    def field(label, value):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        p.add_run(value or "—")

    # Title
    title = doc.add_paragraph()
    tr = title.add_run("Supporting Information")
    tr.bold = True
    tr.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Identification
    idd = D["id"]
    field("Compound", idd["compound_name"])
    field("Code", idd["compound_code"])
    field("Reaction Ref.", idd["reaction_ref"])
    field("Analyst", idd["analyst"])
    field("Date", idd["date"])
    if idd["smiles"]:
        field("SMILES", idd["smiles"])

    # 2. NMR
    active_nmr = [t for t in NMR_TECHNIQUES if D["nmr_checked"].get(t)]
    if active_nmr:
        h("NMR Spectroscopy")
        for t in active_nmr:
            nd = D["nmr"][t]
            meta = f"{t}"
            details = []
            if nd["freq"]:
                details.append(f"{nd['freq']} MHz")
            if nd["solvent"]:
                details.append(nd["solvent"])
            if details:
                meta += f" ({', '.join(details)})"
            if t == "¹H NMR":
                acs = build_acs_1h(nd["rows"])
                p = doc.add_paragraph()
                p.add_run(f"{meta}: ").bold = True
                p.add_run(acs).italic = True
            elif t in ("¹³C NMR", "DEPT-135", "DEPT-90"):
                shifts = [r["delta"] for r in nd["rows"] if r.get("delta")]
                p = doc.add_paragraph()
                p.add_run(f"{meta}: ").bold = True
                p.add_run("δ " + ", ".join(shifts))
            else:
                field(meta, nd["notes"] or "(see attached spectrum)")
            if nd["notes"] and t in ("¹H NMR", "¹³C NMR"):
                doc.add_paragraph(f"Notes: {nd['notes']}")

    # 3. IR
    ir = D["ir"]
    if ir["bands"] or ir["detected_bands"]:
        h("Infrared Spectroscopy (IR)")
        bands_txt = ir["bands"]
        if not bands_txt and ir["detected_bands"]:
            bands_txt = ", ".join(str(b) for b in ir["detected_bands"])
        field(f"IR ({ir['technique']}) ν (cm⁻¹)", bands_txt)

    # 4. GC
    gc = D["gc"]
    if gc["column"] or gc["peaks"]:
        h("Gas Chromatography (GC)")
        if gc["column"]:
            field("Column", gc["column"])
        if gc["temp"]:
            field("Temperature program", gc["temp"])
        if gc["carrier"]:
            field("Carrier gas", gc["carrier"])
        if gc["peaks"]:
            peaks, _ = compute_peak_percentages(gc["peaks"])
            _add_peak_table(doc, peaks)

    # 5. GC-MS
    gcms = D["gcms"]
    if gcms["column"] or gcms["ions"]:
        h("GC-MS")
        if gcms["column"]:
            field("Column", gcms["column"])
        if gcms["rt"]:
            field("Retention time (min)", gcms["rt"])
        if gcms["ions"]:
            field("m/z", gcms["ions"])

    # 6. HPLC
    hplc = D["hplc"]
    if hplc["column"] or hplc["peaks"]:
        h("HPLC")
        if hplc["column"]:
            field("Column", hplc["column"])
        if hplc["mobile"]:
            field("Mobile phase", hplc["mobile"])
        if hplc["uv"]:
            field("UV detection (nm)", hplc["uv"])
        if hplc["ee"]:
            field("ee (%)", hplc["ee"])
        if hplc["peaks"]:
            peaks, _ = compute_peak_percentages(hplc["peaks"])
            _add_peak_table(doc, peaks)

    # 7. HRMS
    hrms = D["hrms"]
    if hrms["formula"] or hrms["calc"]:
        h("High-Resolution Mass Spectrometry (HRMS)")
        p = doc.add_paragraph()
        p.add_run(f"HRMS ({hrms['technique']}) ").bold = True
        txt = f"m/z calcd for {hrms['formula']}"
        if hrms["calc"]:
            txt += f" {hrms['calc']}"
        if hrms["found"]:
            txt += f", found {hrms['found']}"
        p.add_run(txt)

    # 8. Physical properties
    phys_written = False
    if D["mp"]["value"]:
        h("Physical Properties")
        phys_written = True
        lit = f" (lit. {D['mp']['lit']})" if D["mp"]["lit"] else ""
        field("Melting point (°C)", f"{D['mp']['value']}{lit}")
    if D["optrot"]["alpha"]:
        if not phys_written:
            h("Physical Properties")
            phys_written = True
        o = D["optrot"]
        field("[α]D", f"{o['alpha']} (c {o['conc']}, {o['solvent']}, {o['temp']} °C)")
    if D["ri"]["value"]:
        if not phys_written:
            h("Physical Properties")
        field("Refractive index nD", f"{D['ri']['value']} ({D['ri']['temp']} °C)")

    # 9. Elemental analysis
    ea_rows = [r for r in D["ea"]["rows"] if r.get("elem") and (r.get("calc") or r.get("found"))]
    if ea_rows:
        h("Elemental Analysis")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Element", "Calcd. (%)", "Found (%)"
        for r in ea_rows:
            cells = table.add_row().cells
            cells[0].text = r["elem"]
            cells[1].text = str(r["calc"])
            cells[2].text = str(r["found"])

    # 10. Thermal / XRD
    tga = D["tga"]
    if tga["onset"] or tga["detected_onset"]:
        h("Thermogravimetric Analysis (TGA)")
        onset = tga["onset"] or tga["detected_onset"]
        field("Onset of degradation (°C)", onset)
    dsc = D["dsc"]
    if dsc["onset"] or dsc["peak"]:
        h("Differential Scanning Calorimetry (DSC)")
        if dsc["onset"]:
            field("Onset (°C)", dsc["onset"])
        if dsc["peak"]:
            field("Peak (°C)", dsc["peak"])
        if dsc["enthalpy"]:
            field("ΔH (J/g)", dsc["enthalpy"])
    xrd = D["xrd"]
    if xrd["detected_peaks"] or xrd["notes"]:
        h("X-ray Diffraction (XRD)")
        if xrd["detected_peaks"]:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hc = table.rows[0].cells
            hc[0].text, hc[1].text, hc[2].text = "2θ (°)", "Intensity", "d-spacing (Å)"
            for pk in xrd["detected_peaks"]:
                c = table.add_row().cells
                c[0].text = str(pk["two_theta"])
                c[1].text = str(pk["intensity"])
                c[2].text = str(pk["d_spacing"])
        if xrd["notes"]:
            doc.add_paragraph(xrd["notes"])

    # 11. Additional notes
    if D["extra"]["notes"]:
        h("Additional Notes")
        doc.add_paragraph(D["extra"]["notes"])

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_peak_table(doc, peaks):
    """Helper: add a chromatography peak table to the Word document."""
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hc = table.rows[0].cells
    hc[0].text, hc[1].text = "#", "Rt (min)"
    hc[2].text, hc[3].text = "Area", "Area (%)"
    for i, pk in enumerate(peaks, 1):
        c = table.add_row().cells
        c[0].text = str(i)
        c[1].text = str(pk.get("rt", ""))
        c[2].text = str(pk.get("area", ""))
        c[3].text = f"{pk['pct']}" if pk.get("pct") is not None else "—"


# ─────────────────────────────────────────────────────────────────────────────
# Sidebar: Save / Load JSON
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("💾 Session")
    st.caption("Fill in what you have, save, and continue later.")

    # Save: serialize current data to JSON download
    json_str = json.dumps(D, ensure_ascii=False, indent=2)
    fname = (D["id"]["compound_code"] or "SI").replace(" ", "_")
    st.download_button("⬇ Save session (.json)", json_str,
                       file_name=f"{fname}_SI.json", mime="application/json",
                       use_container_width=True)

    # Load: upload a previously saved JSON
    uploaded_json = st.file_uploader("📂 Load session (.json)", type="json")
    if uploaded_json is not None:
        try:
            loaded = json.load(uploaded_json)
            # Merge into a fresh default to tolerate schema additions
            base = default_data()
            base.update(loaded)
            st.session_state.data = base
            st.success("Session loaded — refreshing…")
            st.rerun()
        except Exception as e:
            st.error(f"Could not load file: {e}")

    st.divider()
    # Export Word
    st.header("📄 Export")
    if st.button("Generate Word (.docx)", use_container_width=True):
        buf = build_docx()
        st.download_button("⬇ Download .docx", buf,
                           file_name=f"{fname}_SupportingInfo.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)

    st.divider()
    if st.button("🗑 Clear all fields", use_container_width=True):
        st.session_state.data = default_data()
        st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# Main UI
# ─────────────────────────────────────────────────────────────────────────────
st.title("Supporting Information")
st.caption("LabBook module · One entry per compound")

# ── 1. Identification ────────────────────────────────────────────────────────
with st.expander("1 · Identification (from Reaction Setup)", expanded=True):
    c1, c2, c3 = st.columns(3)
    D["id"]["compound_name"] = c1.text_input("Compound name", D["id"]["compound_name"], key="id_name")
    D["id"]["compound_code"] = c2.text_input("Compound code", D["id"]["compound_code"], key="id_code")
    D["id"]["reaction_ref"] = c3.text_input("Reaction ref.", D["id"]["reaction_ref"], key="id_ref")
    c4, c5, c6 = st.columns(3)
    D["id"]["analyst"] = c4.text_input("Analyst", D["id"]["analyst"], key="id_analyst")
    D["id"]["date"] = c5.text_input("Date", D["id"]["date"], key="id_date")
    D["id"]["smiles"] = c6.text_input("SMILES", D["id"]["smiles"], key="id_smiles")

# ── 2. NMR ───────────────────────────────────────────────────────────────────
with st.expander("2 · NMR Spectroscopy", expanded=True):
    st.write("**Select techniques:**")
    cols = st.columns(len(NMR_TECHNIQUES))
    for i, t in enumerate(NMR_TECHNIQUES):
        D["nmr_checked"][t] = cols[i].checkbox(t, D["nmr_checked"][t], key=f"chk_{t}")

    for t in [x for x in NMR_TECHNIQUES if D["nmr_checked"][x]]:
        nd = D["nmr"][t]
        st.markdown(f"#### {t}")
        m1, m2, m3 = st.columns(3)
        nd["solvent"] = m1.text_input("Solvent", nd["solvent"], key=f"sol_{t}",
                                      placeholder="CDCl₃")
        nd["freq"] = m2.text_input("Frequency (MHz)", nd["freq"], key=f"frq_{t}",
                                   placeholder="400")
        nd["temp"] = m3.text_input("Temp. (°C)", nd["temp"], key=f"tmp_{t}",
                                   placeholder="25")
        pdf = st.file_uploader(f"Attach {t} spectrum (PDF)", type="pdf", key=f"pdf_{t}")
        if pdf:
            nd["pdf"] = pdf.name
        if nd["pdf"]:
            st.caption(f"📎 {nd['pdf']}")

        is_1h = t == "¹H NMR"
        is_13c = t in ("¹³C NMR", "DEPT-135", "DEPT-90")

        if is_1h or is_13c:
            # Editable table via st.data_editor
            if is_1h:
                df = pd.DataFrame(nd["rows"] or [{"delta": "", "mult": "s",
                                                  "J": "", "integ": "", "assign": ""}])
                df = df.reindex(columns=["delta", "mult", "J", "integ", "assign"])
                edited = st.data_editor(
                    df, num_rows="dynamic", key=f"tbl_{t}", use_container_width=True,
                    column_config={
                        "delta": st.column_config.TextColumn("δ (ppm)"),
                        "mult": st.column_config.SelectboxColumn("Mult.", options=MULTIPLICITIES),
                        "J": st.column_config.TextColumn("J (Hz)"),
                        "integ": st.column_config.TextColumn("∫ (nH)"),
                        "assign": st.column_config.TextColumn("Assignment"),
                    })
            else:
                df = pd.DataFrame(nd["rows"] or [{"delta": "", "type": "CH", "assign": ""}])
                df = df.reindex(columns=["delta", "type", "assign"])
                edited = st.data_editor(
                    df, num_rows="dynamic", key=f"tbl_{t}", use_container_width=True,
                    column_config={
                        "delta": st.column_config.TextColumn("δ (ppm)"),
                        "type": st.column_config.SelectboxColumn("Type", options=C_TYPES),
                        "assign": st.column_config.TextColumn("Assignment"),
                    })
            nd["rows"] = edited.fillna("").to_dict("records")

            # Live ACS string for ¹H
            if is_1h:
                acs = build_acs_1h(nd["rows"])
                if acs:
                    st.code(acs, language=None)
        else:
            st.caption("2D spectrum — attach PDF above.")

        nd["notes"] = st.text_area("Notes", nd["notes"], key=f"nts_{t}", height=68)
        st.divider()

# ── 3. IR ────────────────────────────────────────────────────────────────────
with st.expander("3 · Infrared Spectroscopy (IR)"):
    ir = D["ir"]
    c1, c2 = st.columns([1, 3])
    ir["technique"] = c1.selectbox("Technique", ["ATR", "KBr", "Nujol", "Film"],
                                   index=["ATR", "KBr", "Nujol", "Film"].index(ir["technique"]))
    # Raw FTIR .txt upload → auto band detection
    ftir_file = st.file_uploader("Upload FTIR raw data (.txt, 2 columns)",
                                 type="txt", key="ftir_txt")
    if ftir_file:
        content = ftir_file.read().decode("utf-8", errors="ignore")
        x, y = parse_two_column_txt(content)
        if len(x):
            ir["detected_bands"] = detect_ir_bands(x, y)
            fig = go.Figure(go.Scatter(x=x, y=y, mode="lines", line=dict(width=1)))
            fig.update_layout(height=320, margin=dict(l=40, r=20, t=20, b=40),
                              xaxis_title="Wavenumber (cm⁻¹)", yaxis_title="Signal")
            fig.update_xaxes(autorange="reversed")  # IR convention: high→low
            st.plotly_chart(fig, use_container_width=True)
            st.success(f"Detected bands (cm⁻¹): {', '.join(str(b) for b in ir['detected_bands'])}")
    ir["bands"] = st.text_input("Main bands (cm⁻¹) + assignments", ir["bands"], key="ir_bands",
                                placeholder="3420 (O–H), 1735 (C=O ester), 1250 (C–O)")
    ir["notes"] = st.text_area("Notes", ir["notes"], key="ir_notes", height=68)

# ── 4. GC ────────────────────────────────────────────────────────────────────
with st.expander("4 · Gas Chromatography (GC)"):
    gc = D["gc"]
    c1, c2 = st.columns(2)
    gc["column"] = c1.text_input("Column", gc["column"], key="gc_col", placeholder="HP-5, 30 m × 0.25 mm")
    gc["temp"] = c2.text_input("Temperature program", gc["temp"], key="gc_temp", placeholder="60→280 °C, 10 °C/min")
    c3, c4 = st.columns(2)
    gc["carrier"] = c3.text_input("Carrier gas", gc["carrier"], key="gc_carrier", placeholder="He, 1.0 mL/min")
    gc["rt"] = c4.text_input("Rt (min)", gc["rt"], key="gc_rt", placeholder="12.4")
    gc_pdf = st.file_uploader("Attach GC trace (PDF)", type="pdf", key="gc_pdf")
    if gc_pdf:
        gc["pdf"] = gc_pdf.name
    # Peak table with auto area%
    st.markdown("**Peak Table** — Area (%) is computed automatically as area ÷ Σareas")
    df = pd.DataFrame(gc["peaks"] or [{"rt": "", "area": "", "assign": ""}])
    df = df.reindex(columns=["rt", "area", "assign"])
    edited = st.data_editor(df, num_rows="dynamic", key="gc_peaks", use_container_width=True,
                            column_config={"rt": "Rt (min)", "area": "Area",
                                           "assign": "Assignment"})
    gc["peaks"] = edited.fillna("").to_dict("records")
    peaks, total = compute_peak_percentages(gc["peaks"])
    if total > 0:
        st.dataframe(pd.DataFrame([{"Rt": p["rt"], "Area": p["area"],
                                    "Area %": p["pct"], "Assignment": p["assign"]}
                                   for p in peaks if p.get("area")]),
                     use_container_width=True, hide_index=True)
    gc["notes"] = st.text_area("Notes / purity", gc["notes"], key="gc_notes", height=68)

# ── 5. GC-MS ─────────────────────────────────────────────────────────────────
with st.expander("5 · GC-MS"):
    gcms = D["gcms"]
    c1, c2 = st.columns(2)
    gcms["column"] = c1.text_input("Column", gcms["column"], key="gcms_col", placeholder="HP-5MS, 30 m")
    gcms["temp"] = c2.text_input("Temperature program", gcms["temp"], key="gcms_temp", placeholder="60→280 °C")
    c3, c4 = st.columns(2)
    gcms["rt"] = c3.text_input("Rt (min)", gcms["rt"], key="gcms_rt", placeholder="12.4")
    gcms["ions"] = c4.text_input("Main m/z ions", gcms["ions"], key="gcms_ions", placeholder="132 (M⁺), 101, 87")
    gcms_pdf = st.file_uploader("Attach GC-MS (PDF)", type="pdf", key="gcms_pdf")
    if gcms_pdf:
        gcms["pdf"] = gcms_pdf.name
    gcms["notes"] = st.text_area("Notes", gcms["notes"], key="gcms_notes", height=68)

# ── 6. HPLC ──────────────────────────────────────────────────────────────────
with st.expander("6 · HPLC"):
    hplc = D["hplc"]
    c1, c2 = st.columns(2)
    hplc["column"] = c1.text_input("Column", hplc["column"], key="hplc_col", placeholder="Chiralpak IA, 250 × 4.6 mm")
    hplc["mobile"] = c2.text_input("Mobile phase", hplc["mobile"], key="hplc_mobile", placeholder="Hex/iPrOH 95:5")
    c3, c4, c5 = st.columns(3)
    hplc["uv"] = c3.text_input("UV λ (nm)", hplc["uv"], key="hplc_uv", placeholder="254")
    hplc["rt"] = c4.text_input("Rt (min)", hplc["rt"], key="hplc_rt", placeholder="8.2 / 11.5")
    hplc["ee"] = c5.text_input("ee (%)", hplc["ee"], key="hplc_ee", placeholder="98")
    hplc_pdf = st.file_uploader("Attach HPLC trace (PDF)", type="pdf", key="hplc_pdf")
    if hplc_pdf:
        hplc["pdf"] = hplc_pdf.name
    st.markdown("**Peak Table** — Area (%) is computed automatically as area ÷ Σareas")
    df = pd.DataFrame(hplc["peaks"] or [{"rt": "", "area": "", "assign": ""}])
    df = df.reindex(columns=["rt", "area", "assign"])
    edited = st.data_editor(df, num_rows="dynamic", key="hplc_peaks", use_container_width=True,
                            column_config={"rt": "Rt (min)", "area": "Area",
                                           "assign": "Assignment"})
    hplc["peaks"] = edited.fillna("").to_dict("records")
    peaks, total = compute_peak_percentages(hplc["peaks"])
    if total > 0:
        st.dataframe(pd.DataFrame([{"Rt": p["rt"], "Area": p["area"],
                                    "Area %": p["pct"], "Assignment": p["assign"]}
                                   for p in peaks if p.get("area")]),
                     use_container_width=True, hide_index=True)
    hplc["notes"] = st.text_area("Notes", hplc["notes"], key="hplc_notes", height=68)

# ── 7. HRMS ──────────────────────────────────────────────────────────────────
with st.expander("7 · High-Resolution Mass Spectrometry (HRMS)"):
    hrms = D["hrms"]
    ion_opts = ["ESI+", "ESI−", "EI", "APCI+", "APCI−", "MALDI"]
    c1, c2, c3, c4 = st.columns(4)
    hrms["technique"] = c1.selectbox("Ionization", ion_opts,
                                     index=ion_opts.index(hrms["technique"]))
    hrms["formula"] = c2.text_input("Molecular formula", hrms["formula"], key="hrms_formula", placeholder="C₅H₁₀O₃")
    hrms["calc"] = c3.text_input("Calcd. m/z", hrms["calc"], key="hrms_calc", placeholder="119.0708")
    hrms["found"] = c4.text_input("Found m/z", hrms["found"], key="hrms_found", placeholder="119.0711")
    hrms_pdf = st.file_uploader("Attach HRMS (PDF)", type="pdf", key="hrms_pdf")
    if hrms_pdf:
        hrms["pdf"] = hrms_pdf.name
    hrms["notes"] = st.text_area("Notes / adduct", hrms["notes"], key="hrms_notes", height=68)

# ── 8. Physical properties ───────────────────────────────────────────────────
with st.expander("8 · Physical Properties"):
    st.markdown("**Melting Point**")
    c1, c2 = st.columns(2)
    D["mp"]["value"] = c1.text_input("Value (°C)", D["mp"]["value"], key="mp_val", placeholder="82–84")
    D["mp"]["lit"] = c2.text_input("Literature (°C)", D["mp"]["lit"], key="mp_lit", placeholder="83 (ref)")
    st.markdown("**Optical Rotation [α]D**")
    c3, c4, c5, c6 = st.columns(4)
    D["optrot"]["alpha"] = c3.text_input("[α]D", D["optrot"]["alpha"], key="or_alpha", placeholder="+23.4")
    D["optrot"]["conc"] = c4.text_input("c (g/100 mL)", D["optrot"]["conc"], key="or_conc", placeholder="1.0")
    D["optrot"]["solvent"] = c5.text_input("Solvent", D["optrot"]["solvent"], key="or_solvent", placeholder="CHCl₃")
    D["optrot"]["temp"] = c6.text_input("Temp. (°C)", D["optrot"]["temp"], key="or_temp")
    st.markdown("**Refractive Index**")
    c7, c8 = st.columns(2)
    D["ri"]["value"] = c7.text_input("nD", D["ri"]["value"], key="ri_val", placeholder="1.4231")
    D["ri"]["temp"] = c8.text_input("Temp. (°C)", D["ri"]["temp"], key="ri_temp")

# ── 9. Elemental analysis ────────────────────────────────────────────────────
with st.expander("9 · Elemental Analysis"):
    df = pd.DataFrame(D["ea"]["rows"])
    df = df.reindex(columns=["elem", "calc", "found"])
    edited = st.data_editor(df, num_rows="dynamic", key="ea_tbl", use_container_width=True,
                            column_config={"elem": "Element", "calc": "Calcd. (%)",
                                           "found": "Found (%)"})
    D["ea"]["rows"] = edited.fillna("").to_dict("records")
    D["ea"]["notes"] = st.text_area("Notes", D["ea"]["notes"], key="ea_notes", height=68)

# ── 10. Thermal / XRD ────────────────────────────────────────────────────────
with st.expander("10 · Thermal Analysis & XRD"):
    # TGA
    st.markdown("### TGA")
    tga = D["tga"]
    tga_file = st.file_uploader("Upload TGA raw data (.txt, NETZSCH)", type="txt", key="tga_txt")
    if tga_file:
        content = tga_file.read().decode("utf-8", errors="ignore")
        temp, mass, dtg = parse_tga_netzsch(content)
        if len(temp):
            tga["detected_onset"] = detect_tga_onset(temp, mass, dtg)
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=temp, y=mass, name="Mass (%)", line=dict(width=1.5)))
            # Secondary axis for DTG
            fig.add_trace(go.Scatter(x=temp, y=dtg, name="DTG (%/min)",
                                     line=dict(width=1, dash="dot"), yaxis="y2"))
            fig.update_layout(height=340, margin=dict(l=40, r=40, t=20, b=40),
                              xaxis_title="Temperature (°C)",
                              yaxis_title="Mass (%)",
                              yaxis2=dict(title="DTG (%/min)", overlaying="y", side="right"),
                              legend=dict(orientation="h", y=1.1))
            st.plotly_chart(fig, use_container_width=True)
            st.success(f"Detected onset of degradation ≈ {tga['detected_onset']} °C "
                       "(at max mass-loss rate)")
    tga["onset"] = st.text_input("Onset of degradation (°C)", tga["onset"], key="tga_onset",
                                 placeholder=tga["detected_onset"] or "220")
    tga["notes"] = st.text_area("TGA notes", tga["notes"], key="tga_notes", height=68)

    st.divider()
    # DSC
    st.markdown("### DSC")
    dsc = D["dsc"]
    c1, c2, c3 = st.columns(3)
    dsc["onset"] = c1.text_input("Onset (°C)", dsc["onset"], key="dsc_onset", placeholder="78.2")
    dsc["peak"] = c2.text_input("Peak (°C)", dsc["peak"], key="dsc_peak", placeholder="83.1")
    dsc["enthalpy"] = c3.text_input("ΔH (J/g)", dsc["enthalpy"], key="dsc_enthalpy", placeholder="142.3")
    dsc_pdf = st.file_uploader("Attach DSC (PDF)", type="pdf", key="dsc_pdf")
    if dsc_pdf:
        dsc["pdf"] = dsc_pdf.name

    st.divider()
    # XRD
    st.markdown("### X-ray Diffraction (XRD)")
    xrd = D["xrd"]
    xrd["wavelength"] = st.text_input("Wavelength λ (Å)", xrd["wavelength"], key="xrd_wl",
                                      help="Cu Kα = 1.5406 Å (default)")
    drx_file = st.file_uploader("Upload XRD raw data (.txt, Rigaku)", type="txt", key="drx_txt")
    if drx_file:
        content = drx_file.read().decode("utf-8", errors="ignore")
        tt, inten = parse_drx_rigaku(content)
        if len(tt):
            try:
                wl = float(xrd["wavelength"])
            except ValueError:
                wl = 1.5406
            xrd["detected_peaks"] = detect_drx_peaks(tt, inten, wavelength=wl)
            fig = go.Figure(go.Scatter(x=tt, y=inten, mode="lines", line=dict(width=1)))
            # Mark detected peaks
            if xrd["detected_peaks"]:
                fig.add_trace(go.Scatter(
                    x=[p["two_theta"] for p in xrd["detected_peaks"]],
                    y=[p["intensity"] for p in xrd["detected_peaks"]],
                    mode="markers", marker=dict(color="red", size=8), name="Peaks"))
            fig.update_layout(height=340, margin=dict(l=40, r=20, t=20, b=40),
                              xaxis_title="2θ (°)", yaxis_title="Intensity",
                              showlegend=False)
            st.plotly_chart(fig, use_container_width=True)
            if xrd["detected_peaks"]:
                st.dataframe(pd.DataFrame(xrd["detected_peaks"]).rename(
                    columns={"two_theta": "2θ (°)", "intensity": "Intensity",
                             "d_spacing": "d-spacing (Å)"}),
                    use_container_width=True, hide_index=True)
    xrd["notes"] = st.text_area("XRD notes (space group, cell, CCDC no.)",
                                xrd["notes"], key="xrd_notes", height=68)

# ── 11. Additional ───────────────────────────────────────────────────────────
with st.expander("11 · Additional Files & Notes"):
    extra_pdf = st.file_uploader("Attach additional file", key="extra_pdf")
    if extra_pdf:
        D["extra"]["pdf"] = extra_pdf.name
    D["extra"]["notes"] = st.text_area("Additional notes (UV-Vis, CD, references…)",
                                       D["extra"]["notes"], key="extra_notes", height=100)

st.divider()
st.caption("💡 Use the sidebar to save your progress as JSON and resume later, "
           "or export the compiled Supporting Information to Word.")
